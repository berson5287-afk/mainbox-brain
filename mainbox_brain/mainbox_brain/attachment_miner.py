"""
Attachment miner for MaINbox Brain.

The body miner reads what a vendor TYPED in the email.  A large share of real
quotes arrive as attachments -- PDF quote sheets, Excel pricing, occasionally
Word.  This module extracts text from those files and runs it through the SAME
reply fact extractor, so an attached quote produces the same price/stock/lead
facts as a typed one.

Design:
  * Optional dependencies, degrading per-type.  The core brain stays
    stdlib-only; attachment parsing needs libraries and is opt-in:
        py -m pip install pypdf openpyxl python-docx
  * Text-based PDFs and Excel are handled here directly (a real chunk of
    quotes).  SCANNED/image PDFs and image attachments are flagged
    "ocr_needed" -- that is where MaINbox's SmartScan / Ollama-vision OCR
    plugs in; this module does not reimplement OCR.
  * Tables are flattened to "cell | cell | cell" lines so the existing
    line-oriented reply miner can read tabular quotes.

Each extracted fact's evidence line is tagged with the source filename so the
answer layer can show the quote came from an attachment.
"""
from __future__ import annotations
__version__ = "0.51"

import csv
import io
import os
import re
import zipfile
from typing import Optional

import re as _re
from .reply_miner import ReplyFact, _extract_facts, _subject_item, _MONEY, _BARE_PRICED

_HEADER_WORDS = {'item','items','part','parts','number','no','desc','description',
    'price','prices','cost','unit','qty','quantity','stock','lead','leadtime',
    'availability','avail','eta','time','each','net','line','mfg','mfr','catalog',
    'cat','uom','ext','extended','total','amount'}


def _looks_like_header(line: str) -> bool:
    if _re.search(r"\d", line):       # data rows carry part numbers/prices/qty
        return False
    toks = [t for t in _re.findall(r"[A-Za-z]+", line.lower()) if len(t) > 1]
    if len(toks) < 2:
        return False
    hits = sum(1 for t in toks if t in _HEADER_WORDS)
    return hits >= 2 and hits >= len(toks) - 1


def _strip_table_headers(text: str) -> str:
    return "\n".join(ln for ln in text.splitlines() if not _looks_like_header(ln))

# file types we attempt; everything else is ignored
TEXT_PDF = {".pdf"}
SHEET = {".xlsx", ".xlsm"}
SHEET_LEGACY = {".xls"}
DELIMITED = {".csv", ".tsv"}
WORD = {".docx"}
IMAGE = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".gif", ".bmp", ".heic"}
SUPPORTED = TEXT_PDF | SHEET | SHEET_LEGACY | DELIMITED | WORD | IMAGE


def _pip_hint(pkg: str) -> str:
    return f"(install with: py -m pip install {pkg})"


# ---------------------------------------------------------------------------
# Per-type text extraction.  Each returns (text, note).
#   note == "ocr_needed"  -> scanned/image; hand to SmartScan/vision
#   note == "<msg>"        -> could not read (missing dep / error)
#   note == ""             -> text extracted normally
# ---------------------------------------------------------------------------
def _extract_pdf(path: str) -> tuple[str, str]:
    try:
        import pypdf
    except ImportError:
        return "", "pypdf missing " + _pip_hint("pypdf")
    try:
        reader = pypdf.PdfReader(path)
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        text = "\n".join(parts)
        if len(re.sub(r"\s+", "", text)) < 20:
            # no embedded text layer -> almost certainly a scan/image PDF
            return text, "ocr_needed"
        return text, ""
    except Exception as exc:
        return "", f"pdf read error: {exc}"


def _extract_xlsx(path: str) -> tuple[str, str]:
    try:
        import openpyxl
    except ImportError:
        return "", "openpyxl missing " + _pip_hint("openpyxl")
    try:
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        lines: list[str] = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                # keep column POSITIONS (so headers align with data); strip
                # embedded newlines so a cell can't split the row
                cells = ["" if c is None else
                         _re.sub(r"\s+", " ", str(c)).strip() for c in row]
                while cells and not cells[-1]:
                    cells.pop()
                if any(cells):
                    lines.append(" | ".join(cells))
        wb.close()
        return "\n".join(lines), ("" if lines else "empty workbook")
    except Exception as exc:
        return "", f"xlsx read error: {exc}"


def _extract_xls(path: str) -> tuple[str, str]:
    try:
        import xlrd  # type: ignore
    except ImportError:
        return "", "legacy .xls needs xlrd " + _pip_hint("xlrd==1.2.0")
    try:
        book = xlrd.open_workbook(path)
        lines = []
        for sh in book.sheets():
            for r in range(sh.nrows):
                cells = [_re.sub(r"\s+", " ", str(sh.cell_value(r, c))).strip()
                         for c in range(sh.ncols)]
                while cells and not cells[-1]:
                    cells.pop()
                if any(cells):
                    lines.append(" | ".join(cells))
        return "\n".join(lines), ""
    except Exception as exc:
        return "", f"xls read error: {exc}"


def _extract_delimited(path: str) -> tuple[str, str]:
    try:
        with open(path, "r", encoding="utf-8-sig", errors="ignore", newline="") as fh:
            sample = fh.read(4096)
            fh.seek(0)
            delim = "\t" if path.lower().endswith(".tsv") else None
            if delim is None:
                try:
                    delim = csv.Sniffer().sniff(sample, delimiters=",;\t|").delimiter
                except csv.Error:
                    delim = ","
            lines = []
            for row in csv.reader(fh, delimiter=delim):
                cells = [_re.sub(r"\s+", " ", c).strip() for c in row]
                while cells and not cells[-1]:
                    cells.pop()
                if any(cells):
                    lines.append(" | ".join(cells))
        return "\n".join(lines), ""
    except Exception as exc:
        return "", f"csv read error: {exc}"


def _extract_docx(path: str) -> tuple[str, str]:
    # try python-docx, else fall back to stdlib zip+xml (also pulls table cells)
    try:
        import docx  # type: ignore
        d = docx.Document(path)
        lines = [p.text.strip() for p in d.paragraphs if p.text.strip()]
        for tbl in d.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines), ""
    except ImportError:
        pass
    except Exception as exc:
        return "", f"docx read error: {exc}"
    try:
        with zipfile.ZipFile(path) as z:
            xml = z.read("word/document.xml").decode("utf-8", "ignore")
        xml = xml.replace("</w:tr>", "\n").replace("</w:p>", "\n").replace("</w:tc>", " | ")
        text = re.sub(r"<[^>]+>", "", xml)
        text = re.sub(r"[ \t]*\|[ \t]*(?:\n|$)", "\n", text)
        return re.sub(r"\n{2,}", "\n", text).strip(), ""
    except Exception as exc:
        return "", f"docx read error: {exc}"


def extract_text(path: str) -> tuple[str, str]:
    ext = os.path.splitext(path)[1].lower()
    if ext in TEXT_PDF:
        return _extract_pdf(path)
    if ext in SHEET:
        return _extract_xlsx(path)
    if ext in SHEET_LEGACY:
        return _extract_xls(path)
    if ext in DELIMITED:
        return _extract_delimited(path)
    if ext in WORD:
        return _extract_docx(path)
    if ext in IMAGE:
        return "", "ocr_needed"
    return "", f"unsupported type {ext}"


# ---------------------------------------------------------------------------
# Mining
# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------
# Header/column-aware table extraction (structured price sheets & quote tables)
# ---------------------------------------------------------------------------
_HDR_TOKENS = ("part", "item", "cat", "sku", "model", "desc", "price", "cost",
               "net", "list", "qty", "quantity", "pack", "unit", "uom", "per",
               "lead", "eta", "stock", "avail", "number", "amount", "each",
               "ordered", "shipped", "sell", "extended", "ext")
_PART_KEYS = ("part", "item #", "item#", "item no", "itemno", "item", "cat",
              "sku", "model", "number", "stock #", "mfg #", "mpn")
_DESC_KEYS = ("desc",)
_UNIT_KEYS = ("price per", "per", "uom", "unit of", "package style", "pkg style",
              "package", "pkg")
_QTY_KEYS = ("qty", "quantity", "pack", "ordered", "shipped")
_LEAD_KEYS = ("lead", "eta", "ship date", "delivery")
_STOCK_KEYS = ("stock", "avail", "on hand", "status")
# price preference: NET (true cost) > unit price/sell/cost > price; never ext/list/amount(total)
_PRICE_PREF = ("net", "unit price", "sell", "cost", "your price", "price")
_PRICE_AVOID = ("ext", "extended", "amount", "total", "list")
_UNIT_VAL = {"ea": "ea", "ea.": "ea", "each": "ea", "c": "c", "c.": "c",
             "m": "m", "ft": "ft", "lf": "ft", "pr": "pr", "bx": "box",
             "box": "box", "roll": "roll", "reel": "reel", "ct": "ea",
             "jar": "jar", "bag": "bag", "case": "case", "cs": "case",
             "100": "c", "1000": "m"}


def _cols(line: str) -> list[str]:
    return [c.strip() for c in line.split("|")]


def _match_idx(headers: list[str], keys, avoid=()) -> Optional[int]:
    for i, h in enumerate(headers):
        hl = h.lower()
        if any(x in hl for x in avoid):
            continue
        if any(k in hl for k in keys):
            return i
    return None


def _price_idx(headers: list[str]) -> Optional[int]:
    for pref in _PRICE_PREF:
        for i, h in enumerate(headers):
            hl = h.lower()
            if any(x in hl for x in _PRICE_AVOID) and pref not in ("net",):
                continue
            if pref in hl:
                return i
    return None


# columns that are NOT prices even if a price word appears nearby
_NON_PRICE = ("ext", "extended", "amount", "subtotal", "qty", "quantity",
              "pack", "master", "weight", "tensile", "length", "uom",
              "style", "discount", "margin", "%")


def _price_indices(headers: list[str]) -> list[int]:
    """Every column that holds a unit price (so a row with several price tiers
    can be compared). Excludes pack qty, weight, totals, discounts, etc."""
    out = []
    for i, h in enumerate(headers):
        hl = h.lower()
        if any(x in hl for x in _NON_PRICE):
            continue
        if any(k in hl for k in ("net", "price", "cost", "sell", "list",
                                 "your price")):
            out.append(i)
    return out


def _num(cell: str) -> Optional[float]:
    m = _re.search(r"-?\d[\d,]*(?:\.\d+)?", cell.replace("$", ""))
    if not m:
        return None
    try:
        return round(float(m.group(0).replace(",", "")), 4)
    except ValueError:
        return None


def _find_header(lines: list[str]) -> Optional[tuple[int, list[str]]]:
    for i, ln in enumerate(lines[:20]):
        cols = _cols(ln)
        if len(cols) < 3 or _re.search(r"\d\.\d", ln):   # data rows have decimals
            continue
        hits = sum(1 for c in cols if any(t in c.lower() for t in _HDR_TOKENS))
        if hits >= 2:
            return i, cols
    return None


def _table_facts(text: str) -> list[ReplyFact]:
    lines = [l for l in text.splitlines() if l.strip()]
    found = _find_header(lines)
    if not found:
        return []
    hidx, headers = found
    price_cols = _price_indices(headers)
    p_idx = price_cols[0] if price_cols else _price_idx(headers)
    part_idx = _match_idx(headers, _PART_KEYS, avoid=("desc",))
    desc_idx = _match_idx(headers, _DESC_KEYS)
    unit_idx = _match_idx(headers, _UNIT_KEYS, avoid=("price per ", ))
    lead_idx = _match_idx(headers, _LEAD_KEYS)
    stock_idx = _match_idx(headers, _STOCK_KEYS)
    ncols = len(headers)
    if part_idx is None and desc_idx is None:
        return []

    out: list[ReplyFact] = []
    for ln in lines[hidx + 1:]:
        cols = _cols(ln)
        if len(cols) < max(filter(lambda x: x is not None,
                                  [part_idx, desc_idx, p_idx, 0])) + 1:
            continue
        if abs(len(cols) - ncols) > 2:      # not a body row of this table
            continue

        def get(i):
            return cols[i] if (i is not None and i < len(cols)) else ""

        part = get(part_idx)
        desc = get(desc_idx)
        item = (f"{part} {desc}").strip()
        if not item or not _re.search(r"[A-Za-z0-9]", item):
            continue
        # When a row has several price columns (e.g. standard / special /
        # effective), take the LOWEST -- the best available price -- instead of
        # guessing which named column to trust.
        from .reply_miner import _price as _cell_price
        price, unit = None, ""
        candidates = []
        for pi in (price_cols or ([p_idx] if p_idx is not None else [])):
            cell = get(pi)
            pv, pu, _ = _cell_price(cell)         # recovers inline units ($1469/M)
            if pv is None:
                pv = _num(cell)
            if pv is not None and pv > 0:
                candidates.append((pv, pu))
        if candidates:
            price, unit = min(candidates, key=lambda t: t[0])
        if not unit and unit_idx is not None:      # else take unit from its own column
            uv = get(unit_idx).lower().strip().rstrip(".")
            unit = _UNIT_VAL.get(uv, "")
        lead = get(lead_idx).strip()
        stock_raw = get(stock_idx)
        avail, loc = "", ""
        if stock_raw:
            from .reply_miner import _availability
            avail, loc = _availability(stock_raw)
        if price is None and not lead and not avail:
            continue
        conf = 0.45 + (0.25 if price is not None else 0) + (0.1 if unit else 0)
        out.append(ReplyFact(
            source_line=ln[:200], item=item[:120], unit_price=price, unit=unit,
            lead_time=lead[:40], availability=avail, stock_location=loc,
            status=("quoted" if price is not None else "info"),
            confidence=min(conf, 0.95)))
    return out


def looks_like_wire_sheet(text: str) -> bool:
    """A two-panel THHN/THWN + Romex/UF price sheet, identified by its
    distinctive 'LBS/M' weight column header alongside wire types."""
    t = (text or "").upper()
    return "LBS/M" in t and "THHN" in t and ("ROMEX" in t or " UF " in t)


def _wire_sheet_facts(text: str) -> list[ReplyFact]:
    """Parse a two-panel wire price sheet.  Each text line packs a left row and a
    right row as 'SIZE TYPE LBS/M NET'; the NET prices are the only decimals on
    the line (weights are integers), so splitting at each decimal recovers the
    separate products with their correct per-M net price.  This fixes the old
    behavior where two products were mashed into one line and a stray number
    from the Romex column ('12/3') was misread as the price or as a '12' match.
    """
    out: list[ReplyFact] = []
    for raw in text.splitlines():
        line = re.sub(r"\s+", " ", raw.strip())
        if not line or "." not in line:
            continue
        last = 0
        for m in re.finditer(r"[\d,]+\.\d{2}", line):
            seg = line[last:m.start()].replace("$", " ").strip()
            last = m.end()
            try:
                net = float(m.group().replace(",", ""))
            except ValueError:
                continue
            if net <= 0:
                continue
            # seg = "SIZE TYPE... LBS"  (weight is the trailing integer)
            tm = re.match(r"^([0-9][0-9/]*)\s+(.+?)\s+(\d+)$", seg)
            if not tm:
                continue
            size, typ = tm.group(1), tm.group(2).strip()
            if not re.search(r"[A-Za-z]", typ):
                continue
            name = f"{size} {typ}"[:120]
            # source line is THIS product + its price only.  The LBS/M weight is
            # dropped: it's incidental and its digits (e.g. a weight of 12) would
            # otherwise let '16 TFFN STR 12' falsely match a '12 str' search.
            own_line = f"{name} ${net:,.2f}"
            out.append(ReplyFact(source_line=own_line[:200], item=name,
                                 unit_price=net, unit="M", status="quoted",
                                 confidence=0.85))
    return out


def _pdf_layout_lines(path: str) -> Optional[list[str]]:
    """Visual top-to-bottom lines via pdfplumber, which (unlike the default
    extractor) keeps section headers ABOVE their rows -- needed to know that a
    '12' row sits under 'STRANDED' vs 'SOLID'.  Returns None if pdfplumber isn't
    installed, so callers fall back to the flat extractor."""
    try:
        import pdfplumber
    except ImportError:
        return None
    lines: list[str] = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                for ln in page.extract_text_lines():
                    lines.append(ln["text"])
    except Exception:
        return None
    return lines


_WIRE_TYPE_RE = re.compile(
    r"THHN|THWN|XHHW|XHH|RHW|\bUSE\b|USE-2|SOLID|STRANDED|ALUMIN|COPPER|BARE|"
    r"TFFN|MTW|\bSOL\b|\bSTR\b|\bCU\b|\bAL\b", re.I)


def _ctx_from_header(line: str) -> str:
    """Keep just the material/type words from a section header, dropping
    boilerplate ('UL-LISTED', 'BUILDING WIRE', 'PRICE LIST...')."""
    line = re.sub(r"UL[-\s]?LISTED|BUILDING\s+WIRE|PRICE\s+LIST.*|VALID.*|"
                  r"US\s+DOLLARS", " ", line, flags=re.I)
    line = line.replace("/", " ")
    keep = [w for w in line.split() if _WIRE_TYPE_RE.search(w)]
    return " ".join(keep)


def looks_like_sectioned_wire(text: str) -> bool:
    t = (text or "").upper()
    return ("AWG" in t and ("THHN" in t or "XHHW" in t)
            and ("STRANDED" in t or "SOLID" in t
                 or "PRICE/MFT" in t or "PRICE 1FT" in t))


def _wire_table_facts(lines: list[str]) -> list[ReplyFact]:
    """Section-aware price table: rows are 'size $price [$price/MFT]' under a
    header naming the material/type (THHN STRANDED vs SOLID, aluminum, ...).
    The section context is propagated into each product name so a '12' under
    STRANDED is distinct from a '12' under SOLID."""
    out: list[ReplyFact] = []
    ctx = ""
    for raw in lines:
        line = re.sub(r"\s+", " ", raw.strip())
        if not line:
            continue
        dm = re.match(r"^(\d+(?:/\d+)?)\s+\$?\s*([\d,]+\.\d{2,4})"
                      r"(?:\s+\$?\s*([\d,]+\.\d{2}))?\s*$", line)
        if dm and ctx:
            size = dm.group(1)
            per_ft = float(dm.group(2).replace(",", ""))
            per_mft = float(dm.group(3).replace(",", "")) if dm.group(3) else None
            price = per_mft if per_mft is not None else per_ft
            unit = "MFT" if per_mft is not None else "ft"
            name = f"{size} AWG {ctx}".strip()[:120]
            out.append(ReplyFact(source_line=f"{name} ${price:,.2f}/{unit}"[:200],
                                 item=name, unit_price=price, unit=unit,
                                 status="quoted", confidence=0.85))
            continue
        c = _ctx_from_header(line) if _WIRE_TYPE_RE.search(line) else ""
        if c:
            ctx = c
    return out


# -- v0.51: keyword document classification ----------------------------------
# Keywords do what keywords are good at: identifying WHAT KIND of document
# this is and whether it carries money -- never extracting values (proximity
# grabbing is exactly what produced quantity-garbage before v0.50).
_DOC_KIND_RES = [
    ("order_ack",      re.compile(r"\border\s+acknowledg?e?ment\b", re.I)),
    ("invoice",        re.compile(r"\binvoice\b", re.I)),
    ("purchase_order", re.compile(r"\bpurchase\s+order\b|\byour\s+po\b|\bp\.?o\.?\s*(?:number|no\.?|#)", re.I)),
    ("quote",          re.compile(r"\bquotation\b|\bquote\s*(?:number|no\.?|#)|\bquoted?\s+price", re.I)),
    ("price_list",     re.compile(r"\bprice\s+(?:list|sheet|file)s?\b|\bpricing\s+(?:sheet|file)\b", re.I)),
]
# facts from these kinds are transaction prices, not offers
_ORDERED_KINDS = {"order_ack", "invoice", "purchase_order"}


def _doc_kind(text: str) -> str:
    """Classify a document by keywords, first match wins (header order)."""
    head = (text or "")[:3000]
    for kind, rx in _DOC_KIND_RES:
        if rx.search(head):
            return kind
    for kind, rx in _DOC_KIND_RES:          # fall back to the whole doc
        if rx.search(text or ""):
            return kind
    return ""


_BARE_MONEY = re.compile(r"\b\d{1,5}\.\d{2,3}\b")


def _looks_priced(text: str) -> bool:
    """Does this document plainly contain money? (>=2 money-looking figures,
    or a recognized money-document kind). Used to detect priced documents in
    layouts no parser understands yet -- those must be surfaced, not dropped.
    Counts bare decimals too: vendor PDFs rarely bother with $ signs."""
    if _doc_kind(text):
        return True
    if len(_MONEY.findall(text or "")) >= 2:
        return True
    return len(_BARE_MONEY.findall(text or "")) >= 2


def _looks_like_order_ack(text: str) -> bool:
    """v0.50: an Exact-style order acknowledgment (Wesanco-ZSI and others).
    Signature: 'Order Acknowledgment' header + the flattened line-item column
    labels + the 'TOT'-suffixed extended-price convention."""
    t = text or ""
    has_header = re.search(r"\border\s+acknowledg?ment\b", t, re.I) is not None
    has_cols = ("Unit price" in t and "Extended price" in t) or \
               re.search(r"\bItem\s+No\b.*\bQuantity\b", t, re.I) is not None
    has_tot = re.search(r"\d+\.\d{2}\s*TOT\b", t) is not None
    return has_header and (has_cols or has_tot)


# item line: a manufacturer part number on its own line. Broad but anchored --
# has a digit, mostly A-Z0-9 with -/., length 4-24, not a pure quantity/price.
_OA_PART_RE = re.compile(r"^[A-Za-z0-9][A-Za-z0-9\-/.]{3,23}$")
# amount line: "<qty> <UOM> <unit_price> <ext>TOT"  e.g. "50.000 EA  2.240  112.00TOT"
_OA_AMOUNT_RE = re.compile(
    r"^\s*(?P<qty>\d[\d,]*\.?\d*)\s+(?P<uom>[A-Za-z]{1,4})\s+"
    r"(?P<unit>\d[\d,]*\.\d+)\s+(?P<ext>\d[\d,]*\.\d+)\s*TOT\b", re.I)
_OA_NOISE = re.compile(
    r"payment terms|ship via|net \d|pick up|fob\b|promised|shipping|"
    r"customer part|cust #|exact software|sales amount|sales tax|tariff|"
    r"bill to|ship to|formerly|purchase order|e-?mail|^us us$|^\d+$", re.I)


def _to_amt(s: str) -> float | None:
    try:
        return float((s or "").replace(",", ""))
    except (ValueError, AttributeError):
        return None


def _order_ack_facts(text: str) -> list["ReplyFact"]:
    """v0.50: parse Exact-style order-ack line items into priced facts.

    Layout (flattened from a table) repeats per item:
        <PART NUMBER>
        <description...>            (0+ lines)
        <qty> <UOM> <unit> <ext>TOT
    The unit price is the value BEFORE the TOT-suffixed extended price -- the
    old miner grabbed bare numbers and never tied a price to its part, so
    W6128AS4-US ($2.240) was lost. We anchor each amount line to the nearest
    preceding part-number line and keep description for context.
    """
    lines = [ln.rstrip() for ln in (text or "").splitlines()]
    facts: list[ReplyFact] = []
    pending_part = ""
    pending_desc: list[str] = []
    for raw in lines:
        ln = raw.strip()
        if not ln:
            continue
        am = _OA_AMOUNT_RE.match(ln)
        if am and pending_part:
            unit = _to_amt(am.group("unit"))
            ext = _to_amt(am.group("ext"))
            qty = _to_amt(am.group("qty"))
            desc = " ".join(pending_desc).strip()
            item = pending_part + (f" — {desc}" if desc else "")
            facts.append(ReplyFact(
                source_line=f"{pending_part} {qty:g} {am.group('uom')} "
                            f"@ {unit} = {ext}"[:200],
                item=item[:160], unit_price=unit, unit=am.group("uom").upper(),
                ext_price=ext, status="ordered", confidence=0.9,
                direction="cost"))   # v0.51: vendor ack = OUR COST (store vocab)
            pending_part, pending_desc = "", []
            continue
        if _OA_NOISE.search(ln):
            # a noise line breaks a description run but doesn't drop the part
            pending_desc = []
            continue
        if _OA_PART_RE.match(ln) and re.search(r"\d", ln) and " " not in ln:
            # a new part number: previous unpriced part is abandoned (no amount
            # line appeared), this becomes the current one
            pending_part, pending_desc = ln, []
        elif pending_part and len(pending_desc) < 3 and re.search(r"[A-Za-z]", ln):
            pending_desc.append(ln)
    return facts


def mine_attachment(path: str, subject: str = "") -> tuple[list[ReplyFact], str]:
    """Extract reply facts from a single attachment file.

    Returns (facts, note).  note surfaces "ocr_needed" or an error so callers
    can route scans to SmartScan or report missing dependencies.
    """
    text, note = extract_text(path)
    if not text:
        return [], note
    subj_item = _subject_item(subject) if subject else ""
    facts: list[ReplyFact] = []
    # v0.50: order acknowledgments (Exact/Wesanco layout) first -- their
    # flattened line-item tables defeat the generic parser, which grabbed bare
    # quantities and lost the part<->price pairing entirely.
    if _looks_like_order_ack(text):
        facts = _order_ack_facts(text)
    if not facts and looks_like_wire_sheet(text):
        facts = _wire_sheet_facts(text)          # two-panel THHN | Romex sheet
    if not facts and path.lower().endswith(".pdf"):
        layout = _pdf_layout_lines(path)         # layout-ordered (optional dep)
        if layout and looks_like_sectioned_wire("\n".join(layout)):
            facts = _wire_table_facts(layout)    # sectioned STR/SOL price list
    if not facts:
        facts = _table_facts(text)               # structured price/quote tables
        if not facts:
            facts = _extract_facts(_strip_table_headers(text), subject_item=subj_item)
    # a PO number in the document (confirmation/invoice) applies to all its lines
    from .reply_miner import _po_number
    po = _po_number(text)
    tag = os.path.basename(path)
    # v0.51: keyword doc-kind ('purchase order', 'invoice', 'quote'...) tags
    # every fact regardless of which parser produced it, so a price from an
    # invoice reads as a transaction, not an offer.
    kind = _doc_kind(text)
    for f in facts:
        if po and not f.po_number:
            f.po_number = po
        if kind in _ORDERED_KINDS:
            if f.status == "quoted":
                f.status = "ordered"
            if not f.direction:
                f.direction = "cost"
        f.source_line = f"[{tag}] {f.source_line}"
    # v0.51: a priced document that NO parser could read must be surfaced --
    # this is what makes the system universal in practice: unknown layouts
    # announce themselves instead of silently contributing nothing.
    if not facts and _looks_priced(text):
        extra = f"unparsed_priced_doc kind={kind or 'unknown'}"
        note = f"{note}; {extra}" if note else extra
    return facts, note


def mine_record_attachments(rec: dict, base_dir: str) -> tuple[list[ReplyFact], list[str]]:
    """Mine every attachment listed on an export record.

    `rec["attachments"]` is a list of paths relative to `base_dir` (as written
    by export_replies_outlook.py --attachments).  Returns (facts, notes).
    """
    facts: list[ReplyFact] = []
    notes: list[str] = []
    subject = rec.get("subject", "") or ""
    for rel in rec.get("attachments", []) or []:
        full = rel if os.path.isabs(rel) else os.path.join(base_dir, rel)
        if not os.path.exists(full):
            notes.append(f"missing file: {rel}")
            continue
        if os.path.splitext(full)[1].lower() not in SUPPORTED:
            continue
        f, note = mine_attachment(full, subject=subject)
        facts.extend(f)
        if note:
            notes.append(f"{os.path.basename(rel)}: {note}")
    return facts, notes


# ---------------------------------------------------------------------------
# CLI: preview what would be extracted from a file or a folder of attachments
#   py -m mainbox_brain.attachment_miner quote.pdf
#   py -m mainbox_brain.attachment_miner quote_files/
# ---------------------------------------------------------------------------
def _print_file(path: str, subject: str, limit: int) -> tuple[int, int, int]:
    name = os.path.basename(path)
    facts, note = mine_attachment(path, subject=subject)
    if note == "ocr_needed":
        print(f"\n  [SCAN] {name}")
        print(f"         No embedded text -> needs OCR (route to SmartScan / Ollama-vision).")
        return 0, 0, 1
    if note and not facts:
        print(f"\n  [skip] {name}: {note}")
        return 0, 0, 0
    priced = sum(1 for f in facts if f.unit_price is not None)
    print(f"\n  {name}  -- {len(facts)} fact(s), {priced} priced"
          + (f"  (note: {note})" if note else ""))
    for f in facts[:limit]:
        if f.unit_price is not None:
            price = f"${f.unit_price:g}" + (f"/{f.unit}" if f.unit else "")
        else:
            price = f.status
        extra = []
        if f.availability:
            extra.append(f.availability + (f"@{f.stock_location}" if f.stock_location else ""))
        if f.lead_time:
            extra.append("lead=" + f.lead_time)
        if f.eta:
            extra.append("eta=" + f.eta)
        print(f"      {f.item[:48]:<50} {price:<14} {' '.join(extra)}")
    if len(facts) > limit:
        print(f"      ... and {len(facts) - limit} more")
    return len(facts), priced, 0


def main() -> None:
    import argparse
    from pathlib import Path
    ap = argparse.ArgumentParser(
        description="Preview reply facts extracted from PDF/Excel/CSV/Word attachments.")
    ap.add_argument("path", help="an attachment file, or a folder of attachments")
    ap.add_argument("--subject", default="",
                    help="subject line to use as fallback product context")
    ap.add_argument("--limit", type=int, default=20, help="max facts shown per file")
    args = ap.parse_args()

    p = Path(args.path)
    if p.is_file():
        files = [p]
    elif p.is_dir():
        files = sorted(f for f in p.rglob("*")
                       if f.suffix.lower() in SUPPORTED and f.is_file())
    else:
        print(f"Not found: {p}")
        raise SystemExit(1)
    if not files:
        print(f"No supported attachments found in {p} "
              f"(looked for: {', '.join(sorted(SUPPORTED))}).")
        return

    tot_f = tot_p = tot_scan = 0
    for f in files:
        nf, np_, ns = _print_file(str(f), args.subject, args.limit)
        tot_f += nf
        tot_p += np_
        tot_scan += ns
    print(f"\n{len(files)} file(s): {tot_f} fact(s), {tot_p} priced, "
          f"{tot_scan} scanned/needs-OCR.")
    print("To mine attachments into the brain (with vendor attribution), use the "
          "email pipeline:\n  py -m mainbox_brain.reply_corpus received_export.json "
          "--db mainbox.db --attachments quote_files")


if __name__ == "__main__":
    main()
